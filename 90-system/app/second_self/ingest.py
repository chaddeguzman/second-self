from __future__ import annotations

import hashlib
import shutil
from datetime import date
from pathlib import Path

from .paths import SecondSelfPaths


SUPPORTED = {".pdf", ".docx", ".xlsx", ".txt"}


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _extract_pdf(path: Path) -> tuple[str, list[str]]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    if reader.is_encrypted:
        return "", ["encrypted PDF; extraction deferred"]
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(f"## Page {index + 1}\n\n{body}" for index, body in enumerate(pages) if body)
    warnings = [] if text.strip() else ["no extractable text; possible scanned document"]
    return text, warnings


def _extract_docx(path: Path) -> tuple[str, list[str]]:
    from docx import Document

    document = Document(path)
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        blocks.append(f"## Table {table_index}")
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells))
    return "\n\n".join(blocks), []


def _extract_xlsx(path: Path) -> tuple[str, list[str]]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=False)
    output: list[str] = []
    for sheet in workbook.worksheets:
        output.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            if any(value is not None for value in row):
                output.append(" | ".join("" if value is None else str(value) for value in row))
        output.append("")
    return "\n".join(output).strip(), []


def extract(path: Path) -> tuple[str, list[str]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".xlsx":
        return _extract_xlsx(path)
    if suffix == ".txt":
        try:
            return path.read_text(encoding="utf-8-sig"), []
        except UnicodeDecodeError:
            return path.read_text(encoding="cp1252"), ["decoded as Windows-1252"]
    raise ValueError(f"Unsupported import format: {suffix}")


def ingest(paths: SecondSelfPaths, source: Path) -> dict[str, str | list[str] | bool]:
    source = source.resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    if source.suffix.lower() not in SUPPORTED:
        raise ValueError(f"Unsupported import format: {source.suffix or '(none)'}")

    digest = file_hash(source)
    short = digest[:12]
    originals = paths.layer1 / "75-imports" / "originals"
    extracted_dir = paths.layer1 / "75-imports" / "extracted"
    inbox = paths.layer1 / "00-inbox"
    original = originals / f"{digest}{source.suffix.lower()}"
    duplicate = original.exists()
    if not duplicate:
        shutil.copy2(source, original)

    text, warnings = extract(original)
    today = date.today().isoformat()
    extracted = extracted_dir / f"{digest}.md"
    if not extracted.exists():
        warning_text = "\n".join(f"- {warning}" for warning in warnings) or "- None"
        extracted.write_text(
            f"""---
type: import
created: {today}
status: proposed
source: "{original.relative_to(paths.data_root).as_posix()}"
tags: []
projects: []
related: []
---

# Extracted - {source.name}

SHA-256: `{digest}`

## Extraction Warnings

{warning_text}

## Extracted Content

{text}
""",
            encoding="utf-8",
        )

    intake = inbox / f"Import - {source.stem} - {short}.md"
    if not intake.exists():
        intake.write_text(
            f"""---
type: import
created: {today}
status: inbox
source: "{original.relative_to(paths.data_root).as_posix()}"
tags: []
projects: []
related:
  - "{extracted.relative_to(paths.data_root).as_posix()}"
---

# Import - {source.name}

## Provenance

- SHA-256: `{digest}`
- Original filename: `{source.name}`
- Immutable original: [{original.name}](../75-imports/originals/{original.name})
- Extraction: [{extracted.name}](../75-imports/extracted/{extracted.name})

## Proposed Derived Notes

Review the extraction and propose categorized notes. Do not delete either source
artifact after derived notes are approved.
""",
            encoding="utf-8",
        )
    return {
        "sha256": digest,
        "duplicate": duplicate,
        "original": str(original),
        "extracted": str(extracted),
        "intake": str(intake),
        "warnings": warnings,
    }

